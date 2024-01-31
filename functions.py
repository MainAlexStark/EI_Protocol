from docx import Document
from docx.shared import Pt
import docx
import os

def extract_value(text, start_text, end_text):
    if text.find(start_text) != -1:
        start = text.find(start_text) + len(start_text)

        end = text.find(end_text)

        if end_text == '\n':
            value = text[start:]
        else:
            value = text[start:end]

        return value.strip()
def create_template(path:str):

    try:

        doc = Document(path)
        print('Файл успешно открыт')

        # Получение первой таблицы
        table = doc.tables[0]

        # Получение текста из клетки 0 0
        cell = table.cell(0, 0)
        cell_text = cell.text


        # Словарь для переменных
        var = {}

        index = 0
        for paragraph in doc.paragraphs:
            text = paragraph.text

            if extract_value(text, "ПРОТОКОЛ №", "\n"):

                var['номер протокола'] = extract_value(text, "ПРОТОКОЛ №", "\n")

                text = text.replace(extract_value(text, "ПРОТОКОЛ №", "\n"), 'НОМЕР_ПРОТОКОЛА')
                text = text.replace('№', '#', 1)

            if extract_value(text, "СИ –", "№"):
                print(extract_value(text, "СИ –", "№"))
                var['весы'] = extract_value(text, "СИ –", "№")

            if extract_value(text, "№", "\n"):
                var['номер весов'] = extract_value(text, "№", "\n")
                text = text.replace(extract_value(text, "№", "\n"), 'НОМЕР_ВЕСОВ')

            if extract_value(text, "Принадлежащего:", ", ИНН"):
                var['компания'] = extract_value(text, "Принадлежащего:", ", ИНН")
                text = text.replace(extract_value(text, "Принадлежащего:", ", ИНН"), "КОМПАНИЯ")
                text = text.replace(',', '', 1)

            if extract_value(text, "ИНН", ","):
                var['инн'] = extract_value(text, "ИНН", ",")
                text = text.replace(extract_value(text, "ИНН", ","), "НОМЕР_ИНН")

                var['юридический адрес'] = extract_value(text, ",", "\n")
                text = text.replace(extract_value(text, ",", "\n"), "ЮРИДИЧЕСКИЙ_АДРЕС")

            if extract_value(text, "Место поверки:", "\n"):
                var['место поверки'] = extract_value(text, "Место поверки:", "\n")
                text = text.replace(extract_value(text, "Место поверки:", "\n"), "МЕСТО_ПОВЕРКИ")

            if extract_value(text, "ФИФ ОЕИ:", "ОСНОВНЫЕ"):
                var['фиф'] = extract_value(text, "ФИФ ОЕИ:", "\n")

            if extract_value(text, "начале поверки:", "°C"):
                var['температура в начале поверки'] = extract_value(text, "начале поверки:", "°C")
                text = text.replace(extract_value(text, "начале поверки:", "°C"), "ТЕМПЕРАТУРА")

            if extract_value(text, "конце поверки:", "°C"):
                var['температура в конце поверки'] = extract_value(text, "конце поверки:", "°C")
                text = text.replace(extract_value(text, "конце поверки:", "°C"), "ТЕМПЕРАТУРА")

            if extract_value(text, "влажность воздуха:", "%"):
                var['влажность'] = extract_value(text, "влажность воздуха:", "%")
                text = text.replace(extract_value(text, "влажность воздуха:", "%"), "ВЛАЖНОСТЬ")

            if extract_value(text, "давление:", "кПа"):
                var['давление'] = extract_value(text, "давление:", "кПа")
                text = text.replace(extract_value(text, "давление:", "кПа"), "ДАВЛЕНИЕ")

            if 'ЭТАЛОНЫ, применяемые при поверке' in text:
                var['эталоны поверки'] = doc.paragraphs[index + 1].text
                doc.paragraphs[index + 1].text = "ЭТАЛОНЫ_ПОВЕРКИ"

            if extract_value(text, "__", "Дата"):
                var['поверитель'] = extract_value(text, "__", "Дата").replace('_', '')
                text = text.replace(extract_value(text, "__", "Дата").replace('_', ''), "ПОВЕРИТЕЛЬ")

                var['дата поверки'] = extract_value(text, "Дата поверки", " г.")
                text = text.replace(extract_value(text, "Дата поверки", " г."), "ДАТА_ПОВЕРКИ")

            if 'УСЛОВИЯ ПРОВЕДЕНИЯ ПОВЕРКИ' in text:
                doc.paragraphs[index + 1].text = cell_text

            paragraph.text = text
            index += 1

        for paragraph in doc.paragraphs:
            text = paragraph.text

            text = text.replace('#', '№')

            paragraph.text = text

        print(var)

        # Проверяем все ли переменные заполнены
        for i in var:
            if not i: return var[i]

        # Сохраняем файл
        doc.save(f"templates/{var["весы"]} {var['фиф']}.docx")
        print('Файл сохранен')


    except Exception as e:
        print('Ошибка при создании протокола:', e)
        return True

    return False



def make_new_protocol(path,
                      name_scale,
                      num_protocol,
                      num_scale,
                      company,
                      INN,
                      legal_place,
                      ver_place,
                      temperature,
                      humidity,
                      pressure,
                      standards,
                      verificationer,
                      day,
                      month,
                      year,
                      unfit):
    print(path,
          name_scale,
          num_protocol,
          num_scale,
          company,
          INN,
          legal_place,
          ver_place,
          temperature,
          humidity,
          pressure,
          standards,
          verificationer,
          day,
          month,
          year,
          unfit)