from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import docx
import os

from loguru import logger

from . import Excel

EXCEL = Excel


def extract_value(text, start_text, end_text):
    try:
        if text.find(start_text) != -1:
            start = text.find(start_text) + len(start_text)

            end = text.find(end_text)

            if end_text == '\n':
                value = text[start:]
            else:
                value = text[start:end]

            return value.strip()
        
    except Exception as e:
        logger.error(f'Ошибка при получении значения: {e}')
        

class Word():
    
    def open_document(self,path: str):
        logger.debug(f'Open Word Document: {path}')
        
        try:
            doc = Document(path)
            
            logger.debug('Документ Word открыт')
            
            return doc
            
        except Exception as e:
            logger.error(f'Ошибка при открытии Word документа: {e}')
            
            
    def make_new_protocol(self,args):
        
        try:
            logger.debug('start')
            
            num_protocol = args['num_protocol'].split('-',1)[0] + '-' + args['work_place'].split('-')[0].strip() + '-' + args['num_protocol'].split('-',1)[1]

            for key in args.keys():
                argument = args[key]
                logger.debug(f"{key}: {argument}")
                
                
            path_to_word_file = f'app\\templates\\Word\\{args['scale']}.docx'
                
            doc = Word.open_document(self,path_to_word_file)

            index = 0
            for element in doc.element.body:
                # Проверяем, является ли элемент абзацем
                if isinstance(element, docx.oxml.text.paragraph.CT_P):
                    # Получаем абзац
                    paragraph = docx.text.paragraph.Paragraph(element, doc)

                    paragraph.text = paragraph.text.replace('НОМЕР_ПРОТОКОЛА_ПЕРЕМЕННАЯ',num_protocol)\
                        .replace('НОМЕР_ВЕСОВ_ПЕРЕМЕННАЯ',args['num_scale']) \
                        .replace('КОМПАНИЯ_ПЕРЕМЕННАЯ',args['company']) \
                        .replace('НОМЕР_ИНН_ПЕРЕМЕННАЯ', args['INN']) \
                        .replace('ЮРИДИЧЕСКИЙ_АДРЕС_ПЕРЕМЕННАЯ', args['legal_address']) \
                        .replace('МЕСТО_ПОВЕРКИ_ПЕРЕМЕННАЯ', args['inspection_address']) \
                        .replace('ТЕМПЕРАТУРА_ПЕРЕМЕННАЯ', args['temperature']) \
                        .replace('ВЛАЖНОСТЬ_ПЕРЕМЕННАЯ', args['humidity']) \
                        .replace('ДАВЛЕНИЕ_ПЕРЕМЕННАЯ', args['pressure']) \
                        .replace('ЭТАЛОНЫ_ПОВЕРКИ_ПЕРЕМЕННАЯ', args['standarts']) \
                        .replace('ПОВЕРИТЕЛЬ_ПЕРЕМЕННАЯ', args['verificationer']) \
                        .replace('ДАТА_ПОВЕРКИ_ПЕРЕМЕННАЯ', args['inspection_date']) \

                    if 'voltage' in args.keys():
                        paragraph.text = paragraph.text.replace('НАПРЯЖЕНИЕ_ПЕРЕМЕННАЯ', args['voltage']) \
                        .replace('ЧАСТОТА_ПЕРЕМЕННАЯ', args['frequency'])

                    if 'соответствует установленным в описании типа метрологическим требованиям' in paragraph.text and args['unfit'] == True:
                        paragraph.text = paragraph.text.replace('соответствует','несоответствует')\
                        .replace('пригодно','непригодно')
                        
                    if 'Изменение температуры воздуха в помещении в течение ' in paragraph.text:
                        args['change_temperature'] = 'Изменение' + ' ' + extract_value(paragraph.text,'Изменение','°C.') + '°C.'
                        
                    if 'МЕТОДИКА ПОВЕРКИ' in paragraph.text or 'Методика поверки' in paragraph.text:
                        args['method'] = doc.paragraphs[index + 1].text

                    index += 1


            scale = args['scale'].rsplit(' ',1)[0]
            fif = args['scale'].rsplit(' ',1)[1]

            full = f'{args['path']}/{args['num_protocol']} {scale} №{args['num_scale']} ({fif}).docx'

            doc.save(full)
            logger.debug(full)

            if args['use_excel']:
                try:
                    Excel.add_args_to_excel_journal(args=args)
                except Exception as e:
                    logger.error(f'Ошибка при редактировании Excel файла:{e}')
                    
                    
            if args['unfit'] == False:
                    unfit = 'Пригодно'
            else:
                    unfit = 'Непригодно'
                    
                    
            result = ''
            
            result += full + '\n'
            result += args['num_protocol'].split('-',1)[0] + '-' + args['work_place'].split('-')[0].strip() + '-' + args['num_protocol'].split('-',1)[1] + '\n'
            result += args['scale'] + '\n'
            result += args['num_scale'] + '\n'
            result += args['inspection_date'] + '\n'
            result += args['method'] + '\n'
            result += args['temperature'] + '\n'
            result += args['pressure'] + '\n'
            result += args['humidity'] + '\n'
            result += args['standarts'] + '\n'
            result += args['company'] + '\n'
            result += args['verificationer'] + '\n'
            result += args['legal_address'] + '\n'
            result += args['inspection_address'] + '\n'
            result += unfit

            logger.debug('end')

            return result
        
        except Exception as e:
            logger.error(f'Error create protocol: {e}')

            logger.debug('end')

            return e
        
        
        
    def create_template(self,path:str):
        logger.debug('start')

        try:

            doc = Word.open_document(self, path)

            # Получение первой таблицы
            table = doc.tables[0]

            # Получение текста из клетки 0 0
            cell = table.cell(0, 0)
            cell_text = cell.text

            # Получить все таблицы в документе
            tables = doc.tables

            # Удалить таблицу по индексу (нумерация начинается с 0)
            table_index = 0
            table = tables[table_index]

            # Удалить все строки в таблице
            for row in table.rows:
                table._tbl.remove(row._tr)

            # Удалить таблицу из контейнера (документа)
            doc.element.body.remove(table._element)


            # Словарь для переменных
            var = {}

            index = 0
            element_index = 0
            for element in doc.element.body:
                # Проверяем, является ли элемент абзацем
                if isinstance(element, docx.oxml.text.paragraph.CT_P):
                    # Получаем абзац
                    paragraph = docx.text.paragraph.Paragraph(element, doc)

                    if 'УСЛОВИЯ ПРОВЕДЕНИЯ ПОВЕРКИ' in paragraph.text:
                        paragraph.text += '\n' + cell_text
                        paragraph.paragraph_format.alignment = doc.styles['Normal'].paragraph_format.alignment
                        paragraph.paragraph_format.space_after = doc.styles['Normal'].paragraph_format.space_after
                        paragraph.paragraph_format.line_spacing = doc.styles['Normal'].paragraph_format.line_spacing
                        paragraph.paragraph_format.space_before = doc.styles['Normal'].paragraph_format.space_before
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        for run in paragraph.runs:
                            run.font.size = Pt(12)
                            run.font.all_caps = False

                    value = extract_value(paragraph.text, "ПРОТОКОЛ №", "\n")
                    if value:
                        var['номер протокола'] = value

                        paragraph.text = paragraph.text.replace(value, 'НОМЕР_ПРОТОКОЛА_ПЕРЕМЕННАЯ')
                        paragraph.text = paragraph.text.replace('№', '#', 1)

                    value = extract_value(paragraph.text, "периодической поверки", "№")
                    if value:
                        var['весы'] = value.replace('СИ','')

                    value = extract_value(paragraph.text, "СИ -", "№")
                    if value:
                        var['весы'] = value

                    value = extract_value(paragraph.text, "СИ –", "№")
                    if value:
                        var['весы'] = value

                    value = extract_value(paragraph.text, "№", "\n")
                    if value:
                        var['номер весов'] = value
                        paragraph.text = paragraph.text.replace(value, 'НОМЕР_ВЕСОВ_ПЕРЕМЕННАЯ')

                        for temp_paragraph in doc.paragraphs:
                            temp_paragraph.text = temp_paragraph.text.replace('№','#')

                    value = extract_value(paragraph.text, "Принадлежащего:", ", ИНН")
                    if value:
                        var['компания'] = value
                        paragraph.text = paragraph.text.replace(value, "КОМПАНИЯ_ПЕРЕМЕННАЯ")
                        paragraph.text = paragraph.text.replace(',', '', 1)

                    value = extract_value(paragraph.text, "ИНН", ",")
                    if value:
                        var['инн'] = value
                        paragraph.text = paragraph.text.replace(value, "НОМЕР_ИНН_ПЕРЕМЕННАЯ")

                        var['юридический адрес'] = extract_value(paragraph.text, ",", "\n")
                        paragraph.text = paragraph.text.replace(extract_value(paragraph.text, ",", "\n"), "ЮРИДИЧЕСКИЙ_АДРЕС_ПЕРЕМЕННАЯ")

                    value = extract_value(paragraph.text, "Место поверки:", "\n")
                    if value:
                        var['место поверки'] = value
                        paragraph.text = paragraph.text.replace(value, "МЕСТО_ПОВЕРКИ_ПЕРЕМЕННАЯ")

                    value = extract_value(paragraph.text, "ФИФ ОЕИ:", "\n")
                    if value:
                        var['фиф'] = value

                    value = extract_value(paragraph.text, "начале поверки:", "°C")
                    if value:
                        var['температура'] = value
                        paragraph.text = paragraph.text.replace(value, "ТЕМПЕРАТУРА_ПЕРЕМЕННАЯ")

                    value = extract_value(paragraph.text, "конце поверки:", "°C")
                    if value:
                        var['температура'] = value
                        paragraph.text = paragraph.text.replace(value, "ТЕМПЕРАТУРА_ПЕРЕМЕННАЯ")

                    value = extract_value(paragraph.text, "окружающего воздуха:", "°C")
                    if value:
                        var['температура'] = value
                        paragraph.text = paragraph.text.replace(value, "ТЕМПЕРАТУРА_ПЕРЕМЕННАЯ")

                    value = extract_value(paragraph.text, "влажность воздуха:", "%")
                    if value:
                        var['влажность'] = value
                        paragraph.text = paragraph.text.replace(value, "ВЛАЖНОСТЬ_ПЕРЕМЕННАЯ")

                    value = extract_value(paragraph.text, "давление:", "кПа")
                    if value:
                        var['давление'] = value
                        paragraph.text = paragraph.text.replace(value, "ДАВЛЕНИЕ_ПЕРЕМЕННАЯ")

                    value = extract_value(paragraph.text, "Напряжение сети", "В")
                    if value:
                        var['Напряжение'] = value
                        paragraph.text = paragraph.text.replace(value, "НАПРЯЖЕНИЕ_ПЕРЕМЕННАЯ")

                    value = extract_value(paragraph.text, "Частота", "Гц")
                    if value:
                        var['Частота'] = value
                        paragraph.text = paragraph.text.replace(value, "ЧАСТОТА_ПЕРЕМЕННАЯ")

                    if 'ЭТАЛОНЫ, применяемые при поверке' in paragraph.text:
                        j = 0
                        for p in doc.paragraphs:
                            if p.text == 'Методика поверки':
                                break
                            j += 1

                        doc.paragraphs[index + 1].text = "ЭТАЛОНЫ_ПОВЕРКИ_ПЕРЕМЕННАЯ"

                        for par in range(index + 2, j):
                            doc.paragraphs[par].text = None

                    if 'Поверитель' in paragraph.text and '__' in paragraph.text:
                        value = extract_value(paragraph.text, "__", "Дата")
                        if value:
                            var['поверитель'] = value.replace('_', '')
                            paragraph.text = paragraph.text.replace(value.replace('_', ''), "ПОВЕРИТЕЛЬ_ПЕРЕМЕННАЯ")

                            var['дата поверки'] = extract_value(paragraph.text, "Дата поверки", "г.")
                            paragraph.text = paragraph.text.replace(extract_value(paragraph.text, "Дата поверки", "г."), "ДАТА_ПОВЕРКИ_ПЕРЕМЕННАЯ")


                    index += 1
                element_index += 1

            # Loop through paragraphs and tables
            for element in doc.element.body:
                # Проверяем, является ли элемент абзацем
                if isinstance(element, docx.oxml.text.paragraph.CT_P):
                    # Получаем абзац
                    paragraph = docx.text.paragraph.Paragraph(element, doc)

                    # Проверяем наличие символа '#' в тексте абзаца
                    if '#' in paragraph.text:
                        paragraph.text = paragraph.text.replace('#', '№')

            for key in var.keys():
                logger.debug(f'{key}: {var[key]}')

            # Проверяем все ли переменные заполнены
            for i in var:
                if not i: return f'Не удается определить: {var[i]}'

            var['весы'] = var['весы'].replace('/','-').replace('\\','-').replace(':','').replace('*','x').replace('?','').replace('\"',' ').replace('<',' ').replace('>',' ').replace('|',' ')

            # Сохраняем файл

            full = f"app\\templates\\Word\\{var["весы"]} {var['фиф']}.docx"

            doc.save(full)

            logger.debug(f'Файл шаблона сохранен: {full}')


        except Exception as e:
            logger.error(f'Ошибка при создании шаблона: {e}')

            logger.debug('end')

            return e

        return False