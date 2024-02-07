from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import docx
import os

from loguru import logger

# Доделать форматы в Условиях проведения поверки
def extract_value(text, start_text, end_text):
    if text.find(start_text) != -1:
        start = text.find(start_text) + len(start_text)

        end = text.find(end_text)

        if end_text == '\n':
            value = text[start:]
        else:
            value = text[start:end]

        return value.strip()
def create_template(path:str):
    logger.debug('CREATE TEMPLATE')

    file_path = os.path.abspath(__file__)
    main_path = os.path.dirname(os.path.dirname(file_path))

    try:

        doc = Document(path)
        logger.debug(f'File for Template Open ({path})')

        # Получение первой таблицы
        table = doc.tables[0]

        # Получение текста из клетки 0 0
        cell = table.cell(0, 0)
        cell_text = cell.text

        # Получить все таблицы в документе
        tables = doc.tables

        # Удалить таблицу по индексу (нумерация начинается с 0)
        table_index = 0
        table = tables[table_index]

        # Удалить все строки в таблице
        for row in table.rows:
            table._tbl.remove(row._tr)

        # Удалить таблицу из контейнера (документа)
        doc.element.body.remove(table._element)


        # Словарь для переменных
        var = {}

        index = 0
        for element in doc.element.body:
            # Проверяем, является ли элемент абзацем
            if isinstance(element, docx.oxml.text.paragraph.CT_P):
                # Получаем абзац
                paragraph = docx.text.paragraph.Paragraph(element, doc)



                if 'УСЛОВИЯ ПРОВЕДЕНИЯ ПОВЕРКИ' in paragraph.text:
                    paragraph.text += '\n' + cell_text
                    paragraph.paragraph_format.alignment = doc.styles['Normal'].paragraph_format.alignment
                    paragraph.paragraph_format.space_after = doc.styles['Normal'].paragraph_format.space_after
                    paragraph.paragraph_format.line_spacing = doc.styles['Normal'].paragraph_format.line_spacing
                    paragraph.paragraph_format.space_before = doc.styles['Normal'].paragraph_format.space_before
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.font.all_caps = False

                value = extract_value(paragraph.text, "ПРОТОКОЛ №", "\n")
                if value:
                    var['номер протокола'] = value

                    paragraph.text = paragraph.text.replace(value, 'НОМЕР_ПРОТОКОЛА_ПЕРЕМЕННАЯ')
                    paragraph.text = paragraph.text.replace('№', '#', 1)

                value = extract_value(paragraph.text, "периодической поверки", "№")
                if value:
                    var['весы'] = value.replace('СИ','')

                value = extract_value(paragraph.text, "СИ -", "№")
                if value:
                    var['весы'] = value

                value = extract_value(paragraph.text, "СИ –", "№")
                if value:
                    var['весы'] = value

                value = extract_value(paragraph.text, "№", "\n")
                if value:
                    var['номер весов'] = value
                    paragraph.text = paragraph.text.replace(value, 'НОМЕР_ВЕСОВ_ПЕРЕМЕННАЯ')

                    for temp_paragraph in doc.paragraphs:
                        temp_paragraph.text = temp_paragraph.text.replace('№','#')

                value = extract_value(paragraph.text, "Принадлежащего:", ", ИНН")
                if value:
                    var['компания'] = value
                    paragraph.text = paragraph.text.replace(value, "КОМПАНИЯ_ПЕРЕМЕННАЯ")
                    paragraph.text = paragraph.text.replace(',', '', 1)

                value = extract_value(paragraph.text, "ИНН", ",")
                if value:
                    var['инн'] = value
                    paragraph.text = paragraph.text.replace(value, "НОМЕР_ИНН_ПЕРЕМЕННАЯ")

                    var['юридический адрес'] = extract_value(paragraph.text, ",", "\n")
                    paragraph.text = paragraph.text.replace(extract_value(paragraph.text, ",", "\n"), "ЮРИДИЧЕСКИЙ_АДРЕС_ПЕРЕМЕННАЯ")

                value = extract_value(paragraph.text, "Место поверки:", "\n")
                if value:
                    var['место поверки'] = value
                    paragraph.text = paragraph.text.replace(value, "МЕСТО_ПОВЕРКИ_ПЕРЕМЕННАЯ")

                value = extract_value(paragraph.text, "ФИФ ОЕИ:", "\n")
                if value:
                    var['фиф'] = value

                value = extract_value(paragraph.text, "начале поверки:", "°C")
                if value:
                    var['температура'] = value
                    paragraph.text = paragraph.text.replace(value, "ТЕМПЕРАТУРА_ПЕРЕМЕННАЯ")

                value = extract_value(paragraph.text, "конце поверки:", "°C")
                if value:
                    var['температура'] = value
                    paragraph.text = paragraph.text.replace(value, "ТЕМПЕРАТУРА_ПЕРЕМЕННАЯ")

                value = extract_value(paragraph.text, "окружающего воздуха:", "°C")
                if value:
                    var['температура'] = value
                    paragraph.text = paragraph.text.replace(value, "ТЕМПЕРАТУРА_ПЕРЕМЕННАЯ")

                value = extract_value(paragraph.text, "влажность воздуха:", "%")
                if value:
                    var['влажность'] = value
                    paragraph.text = paragraph.text.replace(value, "ВЛАЖНОСТЬ_ПЕРЕМЕННАЯ")

                value = extract_value(paragraph.text, "давление:", "кПа")
                if value:
                    var['давление'] = value
                    paragraph.text = paragraph.text.replace(value, "ДАВЛЕНИЕ_ПЕРЕМЕННАЯ")

                value = extract_value(paragraph.text, "Напряжение сети", "В")
                if value:
                    var['Напряжение'] = value
                    paragraph.text = paragraph.text.replace(value, "НАПРЯЖЕНИЕ_ПЕРЕМЕННАЯ")

                value = extract_value(paragraph.text, "Частота", "Гц")
                if value:
                    var['Частота'] = value
                    paragraph.text = paragraph.text.replace(value, "ЧАСТОТА_ПЕРЕМЕННАЯ")

                if 'ЭТАЛОНЫ, применяемые при поверке' in paragraph.text:
                    j = 0
                    for p in doc.paragraphs:
                        if p.text == 'Методика поверки':
                            break
                        j += 1

                    doc.paragraphs[index + 1].text = "ЭТАЛОНЫ_ПОВЕРКИ_ПЕРЕМЕННАЯ"

                    for par in range(index + 2, j):
                        doc.paragraphs[par].text = None

                if 'Поверитель' in paragraph.text and '__' in paragraph.text:
                    value = extract_value(paragraph.text, "__", "Дата")
                    if value:
                        var['поверитель'] = value.replace('_', '')
                        paragraph.text = paragraph.text.replace(value.replace('_', ''), "ПОВЕРИТЕЛЬ_ПЕРЕМЕННАЯ")

                        var['дата поверки'] = extract_value(paragraph.text, "Дата поверки", "г.")
                        paragraph.text = paragraph.text.replace(extract_value(paragraph.text, "Дата поверки", "г."), "ДАТА_ПОВЕРКИ_ПЕРЕМЕННАЯ")


                index += 1

        # Loop through paragraphs and tables
        for element in doc.element.body:
            # Проверяем, является ли элемент абзацем
            if isinstance(element, docx.oxml.text.paragraph.CT_P):
                # Получаем абзац
                paragraph = docx.text.paragraph.Paragraph(element, doc)

                # Проверяем наличие символа '#' в тексте абзаца
                if '#' in paragraph.text:
                    paragraph.text = paragraph.text.replace('#', '№')

        for key in var.keys():
            logger.debug(f'{key}: {var[key]}')

        # Проверяем все ли переменные заполнены
        for i in var:
            if not i: return f'Не удается определить: {var[i]}'

        var['весы'] = var['весы'].replace('/','-').replace('\\','-').replace(':','').replace('*','x').replace('?','').replace('\"',' ').replace('<',' ').replace('>',' ').replace('|',' ')

        # Сохраняем файл
        doc.save(f"{main_path}/templates/{var["весы"]} {var['фиф']}.docx")
        logger.debug('File was save')


    except Exception as e:
        logger.error(f'Error create tempalte: {e}')
        return e

    return False



def make_new_protocol(args):
    try:
        logger.debug('CREATE PROTOCOL')

        file_path = os.path.abspath(__file__)
        main_path = os.path.dirname(os.path.dirname(file_path))

        for key in args.keys():
            argument = args[key]
            logger.debug(f"{key}: {argument}")

        doc = Document(f'{main_path}\\templates\\{args['scale']}.docx')
        logger.info(f'File for protocol Open: {args['scale']}.docx')

        index = 0
        for element in doc.element.body:
            # Проверяем, является ли элемент абзацем
            if isinstance(element, docx.oxml.text.paragraph.CT_P):
                # Получаем абзац
                paragraph = docx.text.paragraph.Paragraph(element, doc)

                paragraph.text = paragraph.text.replace('НОМЕР_ПРОТОКОЛА_ПЕРЕМЕННАЯ',args['num_protocol'])\
                    .replace('НОМЕР_ВЕСОВ_ПЕРЕМЕННАЯ',args['num_scale']) \
                    .replace('КОМПАНИЯ_ПЕРЕМЕННАЯ',args['company']) \
                    .replace('НОМЕР_ИНН_ПЕРЕМЕННАЯ', args['INN']) \
                    .replace('ЮРИДИЧЕСКИЙ_АДРЕС_ПЕРЕМЕННАЯ', args['legal_address']) \
                    .replace('МЕСТО_ПОВЕРКИ_ПЕРЕМЕННАЯ', args['inspection_address']) \
                    .replace('ТЕМПЕРАТУРА_ПЕРЕМЕННАЯ', args['temperature']) \
                    .replace('ВЛАЖНОСТЬ_ПЕРЕМЕННАЯ', args['humidity']) \
                    .replace('ДАВЛЕНИЕ_ПЕРЕМЕННАЯ', args['pressure']) \
                    .replace('ЭТАЛОНЫ_ПОВЕРКИ_ПЕРЕМЕННАЯ', args['standarts']) \
                    .replace('ПОВЕРИТЕЛЬ_ПЕРЕМЕННАЯ', args['verificationer']) \
                    .replace('ДАТА_ПОВЕРКИ_ПЕРЕМЕННАЯ', args['inspection_date']) \

                if 'voltage' in args.keys():
                    paragraph.text = paragraph.text.replace('НАПРЯЖЕНИЕ_ПЕРЕМЕННАЯ', args['voltage']) \
                    .replace('ЧАСТОТА_ПЕРЕМЕННАЯ', args['frequency'])

                index += 1

        doc.save(f'{args['path']}/{args['scale']}.docx')
        logger.debug(f'Protocol File was Save: {args['path']}/{args['scale']}.docx')





    except Exception as e:
        logger.error(f'Error create protocol: {e}')
        return e